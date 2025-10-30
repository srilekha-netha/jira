import streamlit as st
import requests
from io import BytesIO
from docx import Document
from docx.shared import Pt

from dotenv import load_dotenv
import os

load_dotenv() 

# ---------------------------
# CONFIGURATION
# ---------------------------
st.set_page_config(page_title="AI Jira Defect Writer", layout="wide")
st.title("🧠 AI-Powered Jira Defect Writing Tool")
st.markdown("Paste your user story, issue, and impact area to generate a full Jira defect report with sprint info.")

# ---------------------------
# SIDEBAR INPUTS
# ---------------------------
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")

with st.sidebar:
    st.header("Defect Details")
    sprint_number = st.number_input("🏷️ Sprint Number", min_value=1, step=1)
    module_name = st.text_input("Module Name", placeholder="e.g., ICHRA")
    environment = st.selectbox("🌐 Environment", ["QA", "UAT", "Production"])
    group_id = st.text_input("Group ID", placeholder="Enter Group ID")
    plan_id = st.text_input("Plan ID", placeholder="Enter Plan ID")

# ---------------------------
# MAIN BODY INPUT
# ---------------------------
user_story = st.text_area(
    "Paste User Story / Issue Details",
    placeholder="Describe the issue here...",
    height=200
)

impact_area = st.text_area(
    "Impact Area / Additional Context",
    placeholder="Describe which part of the system is impacted and any other relevant context...",
    height=100
)

generate_btn = st.button("🚀 Generate Jira Defect")

# ---------------------------
# AI GENERATION
# ---------------------------
if generate_btn:
    if not user_story.strip() or not module_name.strip():
        st.warning("⚠️ Please enter Module Name and User Story / Issue details.")
    else:
        # AI prompt tailored for Jira defects with Impact Area
        system_prompt = f"""
        You are an experienced QA Test Engineer. 
        Based on the user story, impact area, and issue, generate a Jira defect in this format exactly:

        TITLE: Sprint {int(sprint_number)} - {module_name} - <short issue title>
        ISSUE DESCRIPTION: <describe the issue in detail including Impact Area>
        STEPS TO REPRODUCE:
        1. Step one
        2. Step two
        ...
        EXPECTED RESULT: <what should happen>
        ACTUAL RESULT: <what actually happens>
        PLAN ID: {plan_id}
        GROUP ID: {group_id}

        Use clear, professional language suitable for Jira. 
        Reference the ICHRA example:
        'As a superuser/Broker user - ICHRA: Unable to update ICHRA settings - getting Invalid request error. 
        Module navigation: Super User -> Group -> Add groups popup -> add ICHRA group -> fill mandatory details -> Save and next -> ICHRA -> Fill settings, contribution and plan details -> Save and Next -> Go back and verify details. Details are not saved. Now enter again and update - Getting 400 - Invalid request body error.'
        """

        user_prompt = f"""
        Module Name: {module_name}
        Environment: {environment}
        Impact Area: {impact_area}
        User Story / Issue Context:
        {user_story}
        """

        with st.spinner("🤖 Generating Jira defect report..."):
            try:
                # API Call to OpenRouter
                url = "https://openrouter.ai/api/v1/chat/completions"
                headers = {
                    "Authorization": f"Bearer {OPENROUTER_API_KEY}",
                    "Content-Type": "application/json",
                }
                payload = {
                    "model": "gpt-4o-mini",
                    "messages": [
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt},
                    ],
                    "temperature": 0.3,
                }

                response = requests.post(url, headers=headers, json=payload)
                response.raise_for_status()
                data = response.json()

                # Safe extraction of AI output
                choices = data.get("choices")
                if not choices or len(choices) == 0:
                    st.error("❌ AI response is empty. Try again.")
                else:
                    message = choices[0].get("message", {})
                    ai_output = message.get("content", "").strip()
                    if not ai_output:
                        st.error("❌ AI returned empty content.")
                    else:
                        # Display in Streamlit (only once)
                        st.success("✅ Generated Jira Defect Report")
                        st.text_area("Jira Defect (Copy Below)", ai_output, height=350, key="jira_defect_text")

                        # Export to Word
                        doc = Document()
                        doc.add_heading(f"Defect Report (Sprint {int(sprint_number)})", level=1)
                        for line in ai_output.splitlines():
                            if line.strip().startswith(("TITLE:", "ISSUE DESCRIPTION:", 
                                                        "STEPS TO REPRODUCE:", 
                                                        "EXPECTED RESULT:", "ACTUAL RESULT:", 
                                                        "PLAN ID:", "GROUP ID:")):
                                p = doc.add_paragraph(line.strip())
                                p.runs[0].bold = True
                                p.runs[0].font.size = Pt(12)
                            elif line.strip():
                                doc.add_paragraph(line.strip())

                        buffer = BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        st.download_button(
                            label="⬇️ Download Jira Defect (Word)",
                            data=buffer,
                            file_name=f"sprint_{int(sprint_number)}_jira_defect.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

            except Exception as e:
                st.error(f"❌ Error generating report: {e}")
